"""
DIYDS Automation API
Endpoints:
  POST /generate-and-publish  — generate PDF + Dropbox + Stripe + GitHub
  POST /publish-guide         — upload existing PDF + Dropbox + Stripe + GitHub
  POST /generate-pdf          — generate PDF from markdown only
  POST /stripe-webhook        — Stripe webhook for automatic PDF email delivery
  GET  /health                — health check
  GET  /token-test            — test all API connections
"""

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
import os, json, base64, time, tempfile, traceback
import requests as req

app = Flask(__name__)
CORS(app)

# ── CREDENTIALS from environment variables ──
DROPBOX_APP_KEY    = os.environ.get('DROPBOX_APP_KEY', '')
DROPBOX_APP_SECRET = os.environ.get('DROPBOX_APP_SECRET', '')
DROPBOX_REFRESH    = os.environ.get('DROPBOX_REFRESH_TOKEN', '')
STRIPE_KEY         = os.environ.get('STRIPE_KEY', '')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
GITHUB_TOKEN       = os.environ.get('GITHUB_TOKEN', '')
GITHUB_REPO        = os.environ.get('GITHUB_REPO', 'diyds/doityadamnself.com')
GITHUB_FILE        = os.environ.get('GITHUB_FILE', 'guides.json')
DROPBOX_FOLDER     = os.environ.get('DROPBOX_FOLDER', '/DIYDS Guides')
SENDGRID_API_KEY   = os.environ.get('SENDGRID_API_KEY', '')
FROM_EMAIL         = os.environ.get('FROM_EMAIL', 'DIYDdonotreply@doityadamnself.com')
FROM_NAME          = os.environ.get('FROM_NAME', 'DoItYaDamnSelf.com')

def get_dropbox_token():
    r = req.post('https://api.dropboxapi.com/oauth2/token', data={
        'grant_type': 'refresh_token',
        'refresh_token': DROPBOX_REFRESH,
        'client_id': DROPBOX_APP_KEY,
        'client_secret': DROPBOX_APP_SECRET,
    })
    r.raise_for_status()
    return r.json()['access_token']

def upload_to_dropbox(file_data, filename):
    token = get_dropbox_token()
    dropbox_path = f'{DROPBOX_FOLDER}/{filename}'

    r = req.post(
        'https://content.dropboxapi.com/2/files/upload',
        headers={
            'Authorization': f'Bearer {token}',
            'Dropbox-API-Arg': json.dumps({'path': dropbox_path, 'mode': 'overwrite', 'autorename': False, 'mute': False}),
            'Content-Type': 'application/octet-stream'
        },
        data=file_data
    )
    if not r.ok:
        raise Exception(f'Dropbox upload failed: {r.status_code} {r.text}')

    token2 = get_dropbox_token()
    r2 = req.post(
        'https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings',
        headers={'Authorization': f'Bearer {token2}', 'Content-Type': 'application/json'},
        json={'path': dropbox_path, 'settings': {'requested_visibility': 'public'}}
    )

    if r2.status_code == 409:
        token3 = get_dropbox_token()
        r3 = req.post(
            'https://api.dropboxapi.com/2/sharing/list_shared_links',
            headers={'Authorization': f'Bearer {token3}', 'Content-Type': 'application/json'},
            json={'path': dropbox_path, 'direct_only': True}
        )
        if not r3.ok:
            raise Exception(f'Dropbox list links failed: {r3.status_code} {r3.text}')
        links = r3.json().get('links', [])
        if not links:
            raise Exception('No shared links found for file')
        url = links[0]['url']
    elif r2.ok:
        url = r2.json()['url']
    else:
        raise Exception(f'Dropbox sharing failed: {r2.status_code} {r2.text}')

    # Clean up URL — strip any existing dl param then add dl=1
    if '?dl=' in url:
        url = url.split('?dl=')[0] + '?dl=1'
    elif '&dl=' in url:
        url = url.split('&dl=')[0] + '&dl=1'
    elif '?' in url:
        url = url + '&dl=1'
    else:
        url = url + '?dl=1'
    return url

def create_stripe_payment_link(title, price_dollars, description=''):
    headers = {'Authorization': f'Bearer {STRIPE_KEY}', 'Content-Type': 'application/x-www-form-urlencoded'}
    r = req.post('https://api.stripe.com/v1/products', headers=headers, data={
        'name': title, 'description': description or title
    })
    r.raise_for_status()
    product_id = r.json()['id']

    r2 = req.post('https://api.stripe.com/v1/prices', headers=headers, data={
        'product': product_id, 'unit_amount': int(float(price_dollars)*100), 'currency': 'usd'
    })
    r2.raise_for_status()
    price_id = r2.json()['id']

    r3 = req.post('https://api.stripe.com/v1/payment_links', headers=headers, data={
        'line_items[0][price]': price_id, 'line_items[0][quantity]': 1,
        'after_completion[type]': 'hosted_confirmation',
        'after_completion[hosted_confirmation][custom_message]': 'Thank you! Check your email for your guide.'
    })
    r3.raise_for_status()
    return r3.json()['url']

def get_github_guides():
    r = req.get(
        f'https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}',
        headers={'Authorization': f'token {GITHUB_TOKEN}', 'Accept': 'application/vnd.github.v3+json'}
    )
    r.raise_for_status()
    data = r.json()
    return json.loads(base64.b64decode(data['content']).decode('utf-8')), data['sha']

def update_github_guides(guides, sha, message):
    content = base64.b64encode(json.dumps(guides, indent=2).encode()).decode()
    r = req.put(
        f'https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}',
        headers={'Authorization': f'token {GITHUB_TOKEN}', 'Accept': 'application/vnd.github.v3+json'},
        json={'message': message, 'content': content, 'sha': sha}
    )
    r.raise_for_status()

def send_guide_email(to_email, customer_name, guide_title, pdf_url, extra_files=None):
    """Send guide delivery email via SendGrid with PDF as attachment"""

    # Download PDF from Dropbox using Dropbox API token instead of public link
    print(f'Downloading PDF from: {pdf_url}')

    # Try multiple download methods
    pdf_data = None

    # Method 1: Direct dl=1 link with browser-like headers
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/pdf,*/*'
        }
        r1 = req.get(pdf_url, allow_redirects=True, timeout=30, headers=headers)
        if r1.ok and len(r1.content) > 1000 and r1.content[:4] == b'%PDF':
            pdf_data = r1.content
            print(f'Method 1 success: {len(pdf_data)} bytes')
    except Exception as e:
        print(f'Method 1 failed: {e}')

    # Method 2: dl.dropboxusercontent.com direct URL
    if not pdf_data:
        try:
            direct_url = pdf_url.replace('www.dropbox.com', 'dl.dropboxusercontent.com')
            direct_url = direct_url.replace('?dl=1', '').replace('?dl=0', '').replace('&dl=1', '').replace('&dl=0', '')
            r2 = req.get(direct_url, allow_redirects=True, timeout=30)
            if r2.ok and len(r2.content) > 1000 and r2.content[:4] == b'%PDF':
                pdf_data = r2.content
                print(f'Method 2 success: {len(pdf_data)} bytes')
        except Exception as e:
            print(f'Method 2 failed: {e}')

    # Method 3: Use Dropbox API to download by path
    if not pdf_data:
        try:
            token = get_dropbox_token()
            # Extract filename from URL
            fname = pdf_url.split('/')[-1].split('?')[0]
            dropbox_path = f'{DROPBOX_FOLDER}/{fname}'
            r3 = req.post(
                'https://content.dropboxapi.com/2/files/download',
                headers={
                    'Authorization': f'Bearer {token}',
                    'Dropbox-API-Arg': json.dumps({'path': dropbox_path})
                },
                timeout=30
            )
            if r3.ok and len(r3.content) > 1000:
                pdf_data = r3.content
                print(f'Method 3 success: {len(pdf_data)} bytes')
        except Exception as e:
            print(f'Method 3 failed: {e}')

    if not pdf_data:
        raise Exception(f'All download methods failed for {pdf_url}')

    pdf_filename = pdf_url.split('/')[-1].split('?')[0]
    if not pdf_filename.endswith('.pdf'):
        pdf_filename += '.pdf'
    pdf_b64 = base64.b64encode(pdf_data).decode()
    print(f'PDF ready: {pdf_filename} ({len(pdf_data)} bytes)')

    # Build attachments list
    attachments = [{
        'content': pdf_b64,
        'filename': pdf_filename,
        'type': 'application/pdf',
        'disposition': 'attachment'
    }]

    # Add extra files if any (e.g. Word template)
    if extra_files:
        for ef in extra_files:
            ef_response = req.get(ef['url'], allow_redirects=True, timeout=30)
            ef_response.raise_for_status()
            ef_data = ef_response.content
            ef_filename = ef['filename']
            ef_b64 = base64.b64encode(ef_data).decode()
            # Determine mime type
            if ef_filename.endswith('.docx'):
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif ef_filename.endswith('.xlsx'):
                mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            else:
                mime = 'application/octet-stream'
            attachments.append({
                'content': ef_b64,
                'filename': ef_filename,
                'type': mime,
                'disposition': 'attachment'
            })

    # HTML email body
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;">
      <div style="background:#1a1a1a;padding:24px;text-align:center;">
        <div style="font-size:22px;font-weight:bold;color:#c5e835;letter-spacing:2px;">DOITYADAMNSELF.COM</div>
        <div style="font-size:12px;color:#888;margin-top:4px;">Your DIY Guide to Getting It Done</div>
      </div>
      <div style="padding:32px;background:#ffffff;">
        <p style="font-size:16px;color:#111;font-weight:bold;">Your guide is attached and ready to use!</p>
        <p style="font-size:14px;color:#444;line-height:1.6;">Thank you for your purchase. Your copy of <strong>{guide_title}</strong> is attached to this email as a PDF.</p>
        <p style="font-size:14px;color:#444;line-height:1.6;">Save it somewhere easy to find — you can reference it as many times as you need.</p>
        <p style="font-size:14px;color:#444;line-height:1.6;">If you have any questions, reply to this email and we'll be happy to help.</p>
        <p style="font-size:14px;color:#111;font-weight:bold;margin-top:24px;">Do It Ya Damn Self! 💪</p>
      </div>
      <div style="background:#c5e835;padding:16px;text-align:center;">
        <div style="font-size:11px;color:#111;">© DoItYaDamnSelf.com — For personal use only. Not legal advice.</div>
      </div>
    </div>
    """

    payload = {
        'personalizations': [{
            'to': [{'email': to_email}],
            'subject': f'Your Guide: {guide_title}'
        }],
        'from': {'email': FROM_EMAIL, 'name': FROM_NAME},
        'content': [{'type': 'text/html', 'value': html_body}],
        'attachments': attachments
    }

    r = req.post(
        'https://api.sendgrid.com/v3/mail/send',
        headers={
            'Authorization': f'Bearer {SENDGRID_API_KEY}',
            'Content-Type': 'application/json'
        },
        json=payload
    )

    if not r.ok:
        raise Exception(f'SendGrid failed: {r.status_code} {r.text}')

    return True

def run_pipeline(pdf_data, filename, title, description, category, price, tags, extra_files=None):
    """Core pipeline — upload to Dropbox and create Stripe link.
    GitHub update is handled separately via /update-github endpoint."""
    dropbox_url = upload_to_dropbox(pdf_data, filename)

    extra_urls = []
    if extra_files:
        for ef_data, ef_filename in extra_files:
            ef_url = upload_to_dropbox(ef_data, ef_filename)
            extra_urls.append({'filename': ef_filename, 'url': ef_url})

    stripe_url = create_stripe_payment_link(title, price, description)

    tag_list = tags if isinstance(tags, list) else [t.strip() for t in tags.split(',') if t.strip()]

    return {
        'success': True,
        'title': title,
        'description': description,
        'category': category,
        'price': str(price),
        'stripe': stripe_url,
        'pdf': dropbox_url,
        'tags': tag_list,
        'extra_files': extra_urls,
        'message': f'PDF uploaded. Guide ready to publish.'
    }

# ── ROUTES ──

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'DIYDS Automation API'})

@app.route('/token-test', methods=['GET'])
def token_test():
    results = {}
    try:
        get_dropbox_token()
        results['dropbox'] = 'ok'
    except Exception as e:
        results['dropbox'] = str(e)
    try:
        guides, _ = get_github_guides()
        results['github'] = f'ok — {len(guides)} guides'
    except Exception as e:
        results['github'] = str(e)
    try:
        r = req.get('https://api.stripe.com/v1/products?limit=1',
                    headers={'Authorization': f'Bearer {STRIPE_KEY}'})
        results['stripe'] = f'ok — status {r.status_code}'
    except Exception as e:
        results['stripe'] = str(e)
    results['sendgrid'] = 'configured' if SENDGRID_API_KEY else 'missing key'
    return jsonify(results)

@app.route('/generate-pdf', methods=['POST'])
def generate_pdf():
    pdf_path = None
    try:
        data = request.get_json()
        if not data or 'markdown' not in data:
            return jsonify({'error': 'Missing markdown field'}), 400

        import sys
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from md2pdf import convert_md_to_pdf

        tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf_path = tmp.name
        tmp.close()

        convert_md_to_pdf(data['markdown'], pdf_path)

        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            return jsonify({'error': 'PDF generation produced empty file'}), 500

        filename = ''.join(c for c in data.get('title','guide').lower().replace(' ','-')
                          if c.isalnum() or c=='-').strip('-') + '.pdf'

        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        os.unlink(pdf_path)

        return Response(pdf_data, mimetype='application/pdf',
                       headers={'Content-Disposition': f'attachment; filename="{filename}"',
                                'Content-Length': str(len(pdf_data))})
    except Exception as e:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        return jsonify({'error': str(e)}), 500

@app.route('/generate-and-publish', methods=['POST'])
def generate_and_publish():
    pdf_path = None
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Missing request body'}), 400

        markdown    = data.get('markdown', '')
        title       = data.get('title', '')
        description = data.get('description', '')
        category    = data.get('category', 'money-credit')
        price       = data.get('price', '2.99')
        tags        = data.get('tags', [])
        custom_filename = data.get('filename', '')

        if not markdown:
            return jsonify({'error': 'Missing markdown field'}), 400
        if not title:
            return jsonify({'error': 'Missing title field'}), 400

        import sys
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from md2pdf import convert_md_to_pdf

        tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf_path = tmp.name
        tmp.close()

        convert_md_to_pdf(markdown, pdf_path)

        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            return jsonify({'error': 'PDF generation produced empty file'}), 500

        if custom_filename:
            filename = custom_filename
        else:
            filename = ''.join(c for c in title.lower().replace(' ','-')
                              if c.isalnum() or c=='-').strip('-') + '.pdf'

        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        os.unlink(pdf_path)
        pdf_path = None

        result = run_pipeline(pdf_data, filename, title, description, category, price, tags)
        return jsonify(result)

    except Exception as e:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/publish-guide', methods=['POST'])
def publish_guide():
    try:
        title       = request.form.get('title', '')
        description = request.form.get('description', '')
        category    = request.form.get('category', 'money-credit')
        price       = request.form.get('price', '2.99')
        tags        = [t.strip() for t in request.form.get('tags', '').split(',') if t.strip()]

        if not title:
            return jsonify({'error': 'Missing title'}), 400

        pdf_file = request.files.get('pdf')
        if not pdf_file:
            return jsonify({'error': 'Missing PDF file'}), 400

        pdf_data = pdf_file.read()
        filename = pdf_file.filename or f"{title.lower().replace(' ','-')}.pdf"

        extra_files = []
        for ef in request.files.getlist('extra_files'):
            extra_files.append((ef.read(), ef.filename))

        result = run_pipeline(pdf_data, filename, title, description, category, price, tags, extra_files)
        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/generate-and-publish-with-extras', methods=['POST'])
def generate_and_publish_with_extras():
    """Generate PDF from markdown + upload extra files (Word template, spreadsheet, etc.)"""
    pdf_path = None
    try:
        markdown    = request.form.get('markdown', '')
        title       = request.form.get('title', '')
        description = request.form.get('description', '')
        category    = request.form.get('category', 'money-credit')
        price       = request.form.get('price', '2.99')
        tags_raw    = request.form.get('tags', '')
        tags        = [t.strip() for t in tags_raw.split(',') if t.strip()]

        if not markdown:
            return jsonify({'error': 'Missing markdown field'}), 400
        if not title:
            return jsonify({'error': 'Missing title field'}), 400

        import sys
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from md2pdf import convert_md_to_pdf

        tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf_path = tmp.name
        tmp.close()

        convert_md_to_pdf(markdown, pdf_path)

        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            return jsonify({'error': 'PDF generation produced empty file'}), 500

        filename = ''.join(c for c in title.lower().replace(' ','-')
                          if c.isalnum() or c=='-').strip('-') + '.pdf'

        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        os.unlink(pdf_path)
        pdf_path = None

        # Handle extra files
        extra_files = []
        for ef in request.files.getlist('extra_files'):
            extra_files.append((ef.read(), ef.filename))

        result = run_pipeline(pdf_data, filename, title, description, category, price, tags, extra_files)
        return jsonify(result)

    except Exception as e:
        if pdf_path and os.path.exists(pdf_path):
            os.unlink(pdf_path)
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500




@app.route('/generate-guide', methods=['POST'])
def generate_guide():
    """Call Claude API to generate a guide with optional attachment"""
    try:
        data = request.get_json()
        if not data or 'request' not in data:
            return jsonify({'error': 'Missing request field'}), 400

        topic = data['request']
        depth = data.get('depth', 'standard')
        extra_request = data.get('extra_request', '').strip()

        depth_instructions = {
            'quick': 'Guide length: CONCISE — 3 to 5 pages. Cover only the essential steps. Skip background info, keep instructions brief and actionable.',
            'standard': 'Guide length: COMPREHENSIVE — 7 to 10 pages. Cover all steps thoroughly with context, tips, and common mistakes.',
            'detailed': 'Guide length: EXHAUSTIVE — 12 to 15 pages. Cover every detail, edge case, variation, and scenario. Include extensive tips, warnings, and examples.'
        }
        depth_note = depth_instructions.get(depth, depth_instructions['standard'])

        # Determine if attachment is requested
        attachment_requested = bool(extra_request)
        attachment_type = 'none'
        if extra_request:
            er_lower = extra_request.lower()
            if any(w in er_lower for w in ['excel', 'spreadsheet', 'xlsx', 'tracker', 'budget']):
                attachment_type = 'excel'
            elif any(w in er_lower for w in ['word', 'doc', 'docx', 'template', 'contract', 'letter', 'agreement', 'nda']):
                attachment_type = 'word'
            else:
                attachment_type = 'word'  # default to word for generic template requests

        import anthropic, json as _json, re as _re
        client = anthropic.Anthropic(api_key=os.environ.get('ANTHROPIC_API_KEY', ''))

        # Build attachment instruction for the guide
        attachment_note = ''
        if extra_request:
            attachment_note = f'''\n\nADDITIONAL REQUEST: {extra_request}
If this involves creating a template or attachment document, acknowledge it in the guide text with a note like "This guide includes a [template name] — see the attached file."
The actual template will be generated separately.'''

        # Generate the guide
        message = client.messages.create(
            model='claude-sonnet-4-20250514',
            max_tokens=12000,
            system=f'''You are an expert at creating practical DIY guides for DoItYaDamnSelf.com.
Respond with a JSON object only (no markdown, no backticks) containing:
{{
  "title": "exact guide title starting with How to",
  "markdown": "complete guide markdown content",
  "price": "2.99, 3.99, or 4.99 based on complexity",
  "category": "one of: money-credit, life-legal, business-organizations, paperwork-filings, transportation-logistics",
  "description": "2-sentence description under 160 characters",
  "tags": ["tag1", "tag2", "tag3", "tag4", "tag5"],
  "attachment_name": "suggested filename for attachment if requested, e.g. nda-template.docx or budget-tracker.xlsx, otherwise empty string"
}}

{depth_note}

Price guide: $2.99 = simple, $3.99 = moderate, $4.99 = complex.

CRITICAL LINK RULE: Every website, form, agency, or tool mentioned ANYWHERE must be a clickable [text](url) link. Never write plain URLs or unlinked website names.

Correct: Visit [AnnualCreditReport.com](https://www.annualcreditreport.com) for free reports
Wrong: Visit AnnualCreditReport.com for free reports{attachment_note}''',
            messages=[{{'role': 'user', 'content': f'Create a DIY guide for: {{topic}}'}}]
        )

        text = message.content[0].text
        try:
            parsed = _json.loads(text)
        except:
            match = _re.search(r'\{{[\s\S]*\}}', text)
            if match:
                parsed = _json.loads(match.group(0))
            else:
                raise Exception('Could not parse AI response')

        attachment_url = None
        attachment_name = parsed.get('attachment_name', '') or ''

        # Generate attachment if requested
        if extra_request and attachment_type != 'none':
            try:
                if attachment_type == 'excel':
                    att_message = client.messages.create(
                        model='claude-sonnet-4-20250514',
                        max_tokens=4000,
                        system='''Create an Excel template. Respond with JSON only:
{
  "sheets": [{"name": "Sheet name", "headers": ["Col1", "Col2"], "rows": [["example", "data"]]}]
}
Make it practical with real example data rows.''',
                        messages=[{{'role': 'user', 'content': f'Create an Excel template for: {{parsed["title"]}}. Additional request: {{extra_request}}'}}]
                    )
                    att_data = _json.loads(att_message.content[0].text)

                    import openpyxl
                    from openpyxl.styles import Font, PatternFill, Alignment
                    import tempfile

                    wb = openpyxl.Workbook()
                    wb.remove(wb.active)

                    for sheet_data in att_data.get('sheets', []):
                        ws = wb.create_sheet(sheet_data['name'])
                        headers = sheet_data.get('headers', [])
                        for ci, header in enumerate(headers, 1):
                            cell = ws.cell(row=1, column=ci, value=header)
                            cell.font = Font(bold=True, color='111111')
                            cell.fill = PatternFill('solid', fgColor='C5E835')
                            cell.alignment = Alignment(horizontal='center')
                            ws.column_dimensions[cell.column_letter].width = max(15, len(str(header)) + 4)
                        for ri, row in enumerate(sheet_data.get('rows', []), 2):
                            for ci, val in enumerate(row, 1):
                                ws.cell(row=ri, column=ci, value=val)

                    tmp = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
                    xl_path = tmp.name
                    tmp.close()
                    wb.save(xl_path)
                    with open(xl_path, 'rb') as f:
                        att_bytes = f.read()
                    os.unlink(xl_path)

                elif attachment_type == 'word':
                    att_message = client.messages.create(
                        model='claude-sonnet-4-20250514',
                        max_tokens=6000,
                        system='''Create a professional Word document template. Respond with JSON only:
{
  "title": "Document title",
  "sections": [{"heading": "Section heading", "content": "template text with [PLACEHOLDER] for fillable parts"}]
}
Make it complete, professional, and ready to use with clear placeholders.''',
                        messages=[{{'role': 'user', 'content': f'Create a Word template for: {{parsed["title"]}}. Additional request: {{extra_request}}'}}]
                    )
                    att_text = att_message.content[0].text
                    try:
                        att_data = _json.loads(att_text)
                    except:
                        match = _re.search(r'\{{[\s\S]*\}}', att_text)
                        att_data = _json.loads(match.group(0)) if match else {'title': parsed['title'], 'sections': []}

                    from docx import Document
                    from docx.shared import Pt, RGBColor, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    import tempfile

                    doc = Document()

                    # Header
                    header = doc.sections[0].header
                    header_para = header.paragraphs[0]
                    header_para.text = 'DoItYaDamnSelf.com'
                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Title
                    title_para = doc.add_heading(att_data.get('title', parsed['title']), 0)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc.add_paragraph()

                    for section in att_data.get('sections', []):
                        doc.add_heading(section.get('heading', ''), level=1)
                        p = doc.add_paragraph(section.get('content', ''))

                    # Footer
                    footer = doc.sections[0].footer
                    footer_para = footer.paragraphs[0]
                    footer_para.text = '© DoItYaDamnSelf.com — For personal use only. Not legal advice.'
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
                    docx_path = tmp.name
                    tmp.close()
                    doc.save(docx_path)
                    with open(docx_path, 'rb') as f:
                        att_bytes = f.read()
                    os.unlink(docx_path)

                # Upload to Dropbox
                if not attachment_name:
                    safe_title = parsed['title'].lower().replace(' ', '-').replace('(', '').replace(')', '').replace("'", '')
                    ext = '.xlsx' if attachment_type == 'excel' else '.docx'
                    attachment_name = safe_title + '-template' + ext

                attachment_url = upload_to_dropbox(att_bytes, attachment_name)
                print(f'✓ Attachment uploaded: {{attachment_name}}')

            except Exception as e:
                print(f'Attachment generation failed: {{e}}')
                traceback.print_exc()

        return jsonify({
            'success': True,
            'title': parsed.get('title', ''),
            'markdown': parsed.get('markdown', ''),
            'price': parsed.get('price', '2.99'),
            'category': parsed.get('category', 'life-legal'),
            'description': parsed.get('description', ''),
            'tags': parsed.get('tags', []),
            'attachment_type': attachment_type if extra_request else 'none',
            'attachment_name': attachment_name,
            'attachment_url': attachment_url,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/get-guides', methods=['GET'])
def get_guides():
    """Get current guides from GitHub"""
    try:
        guides, sha = get_github_guides()
        return jsonify({'success': True, 'guides': guides, 'count': len(guides)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/update-github', methods=['POST'])
def update_github():
    """Update guides.json on GitHub via Railway (keeps token secure)"""
    try:
        data = request.get_json()
        if not data or 'guides' not in data:
            return jsonify({'error': 'Missing guides data'}), 400

        guides = data['guides']
        message = data.get('message', 'Update guides')

        # Get current SHA
        r = req.get(
            f'https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}',
            headers={'Authorization': f'token {GITHUB_TOKEN}', 'Accept': 'application/vnd.github.v3+json'}
        )
        r.raise_for_status()
        sha = r.json()['sha']

        # Update file
        content = base64.b64encode(json.dumps(guides, indent=2).encode()).decode()
        r2 = req.put(
            f'https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}',
            headers={'Authorization': f'token {GITHUB_TOKEN}', 'Accept': 'application/vnd.github.v3+json'},
            json={'message': message, 'content': content, 'sha': sha}
        )
        r2.raise_for_status()
        return jsonify({'success': True, 'message': 'GitHub updated'})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/stripe-webhook', methods=['POST'])
def stripe_webhook():
    """
    Stripe sends a webhook here when a purchase is completed.
    We look up which guide was bought and email the PDF to the customer.
    """
    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature', '')

    # Verify webhook signature if secret is set
    if STRIPE_WEBHOOK_SECRET:
        try:
            import hmac, hashlib
            timestamp = sig_header.split('t=')[1].split(',')[0]
            signatures = [s.split('v1=')[1] for s in sig_header.split(',') if s.startswith('v1=')]
            signed_payload = f'{timestamp}.{payload.decode()}'
            expected = hmac.new(
                STRIPE_WEBHOOK_SECRET.encode(),
                signed_payload.encode(),
                hashlib.sha256
            ).hexdigest()
            if expected not in signatures:
                return jsonify({'error': 'Invalid signature'}), 400
        except Exception as e:
            print(f'Webhook signature error: {e}')

    try:
        event = request.get_json()
        event_type = event.get('type', '')

        if event_type == 'checkout.session.completed':
            session = event['data']['object']

            # Get customer email
            customer_email = session.get('customer_details', {}).get('email', '')
            customer_name = session.get('customer_details', {}).get('name', 'there')

            if not customer_email:
                return jsonify({'status': 'no email'}), 200

            # Get line items to find which guide was purchased
            session_id = session['id']
            r = req.get(
                f'https://api.stripe.com/v1/checkout/sessions/{session_id}/line_items',
                headers={'Authorization': f'Bearer {STRIPE_KEY}'}
            )
            r.raise_for_status()
            line_items = r.json().get('data', [])

            if not line_items:
                return jsonify({'status': 'no line items'}), 200

            # Get the product name from the first line item
            product_name = line_items[0].get('description', '')

            # Look up the guide in guides.json by matching title
            guides, _ = get_github_guides()
            matched_guide = None
            for guide in guides:
                if guide.get('title', '').lower() in product_name.lower() or \
                   product_name.lower() in guide.get('title', '').lower():
                    matched_guide = guide
                    break

            if not matched_guide:
                print(f'No guide matched for product: {product_name}')
                return jsonify({'status': 'no guide matched', 'product': product_name}), 200

            # Send the email with PDF attached
            extra_files = matched_guide.get('extraFiles', [])
            send_guide_email(
                to_email=customer_email,
                customer_name=customer_name or 'there',
                guide_title=matched_guide['title'],
                pdf_url=matched_guide['pdf'],
                extra_files=extra_files
            )

            print(f'✓ Email sent to {customer_email} for: {matched_guide["title"]}')
            return jsonify({'status': 'email sent', 'to': customer_email, 'guide': matched_guide['title']}), 200

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

    return jsonify({'status': 'ok'}), 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
